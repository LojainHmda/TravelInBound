"""
Voucher/Trip Plan Word Document Generator - Windows of Jordan Format
Generates professional Word documents matching the Windows of Jordan voucher format
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import os


class VoucherTripPlanGenerator:
    """Generate Word documents for trip vouchers in Windows of Jordan format"""
    
    def __init__(self):
        self.doc = Document()
        self._setup_styles()
        # Set narrow margins
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Cm(1.27)
            section.bottom_margin = Cm(1.27)
            section.left_margin = Cm(1.27)
            section.right_margin = Cm(1.27)
    
    def _setup_styles(self):
        """Setup document styles"""
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(9)
    
    def _set_cell_shading(self, cell, color):
        """Set background color for table cell"""
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), color)
        cell._element.get_or_add_tcPr().append(shading_elm)
    
    def _add_header_row(self, tour_file, company_name="Windows of Jordan"):
        """Add header with tour file and company name"""
        # Create header table
        table = self.doc.add_table(rows=1, cols=2)
        table.autofit = False
        table.allow_autofit = False
        
        # Set column widths
        table.columns[0].width = Inches(3.0)
        table.columns[1].width = Inches(4.5)
        
        # Tour File cell
        cell = table.rows[0].cells[0]
        p = cell.paragraphs[0]
        run = p.add_run(f'Tour File : {tour_file}')
        run.font.size = Pt(12)
        run.font.bold = True
        
        # Company name cell
        cell = table.rows[0].cells[1]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(company_name)
        run.font.size = Pt(14)
        run.font.bold = True
    
    def _add_tour_details_table(self, tour_data):
        """Add tour details table with group info"""
        # Create table
        table = self.doc.add_table(rows=4, cols=6)
        table.style = 'Table Grid'
        table.autofit = False
        
        # Row 1: Group Name and Visa
        table.rows[0].cells[0].text = 'Group Name:'
        table.rows[0].cells[1].merge(table.rows[0].cells[2])
        table.rows[0].cells[1].text = tour_data.get('guest_name', '')
        table.rows[0].cells[3].text = 'Visa - Free'
        table.rows[0].cells[3].merge(table.rows[0].cells[5])
        
        # Row 2: Nationality and Last Modified
        table.rows[1].cells[0].text = 'Nationality:'
        table.rows[1].cells[1].merge(table.rows[1].cells[2])
        table.rows[1].cells[1].text = tour_data.get('nationality', '')
        table.rows[1].cells[3].text = f"Last Modified: {datetime.now().strftime('%d-%m-%Y %H:%M')}"
        table.rows[1].cells[3].merge(table.rows[1].cells[5])
        
        # Row 3: Agent Name, Pax, Contact Name
        table.rows[2].cells[0].text = 'Agent Name:'
        table.rows[2].cells[1].merge(table.rows[2].cells[2])
        table.rows[2].cells[1].text = tour_data.get('agent_ref', '')
        table.rows[2].cells[3].text = 'Pax :'
        table.rows[2].cells[4].text = str(tour_data.get('pax', ''))
        table.rows[2].cells[5].text = 'Contact Name'
        
        # Row 4: Notes
        table.rows[3].cells[0].text = 'Notes :'
        table.rows[3].cells[1].merge(table.rows[3].cells[5])
        table.rows[3].cells[1].text = tour_data.get('notes', '')
        
        self.doc.add_paragraph()
    
    def _add_arrivals_departures_table(self, arrivals_data):
        """Add Arrivals and Departures table"""
        # Section heading
        heading = self.doc.add_paragraph()
        run = heading.add_run('Arrivals and Departures')
        run.bold = True
        run.font.size = Pt(11)
        run.italic = True
        
        # Create table
        table = self.doc.add_table(rows=1, cols=9)
        table.style = 'Table Grid'
        
        # Header row
        headers = ['Date', 'Border', 'Drop Point', 'Pax', 'Carrier', 'Flight #', 'Time', 'Note']
        header_row = table.rows[0]
        for idx, header_text in enumerate(headers):
            cell = header_row.cells[idx] if idx < 8 else header_row.cells[8]
            cell.text = header_text
            self._set_cell_shading(cell, 'FFFF00')  # Yellow
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add data rows
        for arrival in arrivals_data:
            row = table.add_row()
            row.cells[0].text = arrival.get('date', '')
            row.cells[1].text = arrival.get('border', '')
            row.cells[2].text = arrival.get('drop_point', '')
            row.cells[3].text = str(arrival.get('pax', ''))
            row.cells[4].text = arrival.get('carrier', '')
            row.cells[5].text = arrival.get('flight', '')
            row.cells[6].text = arrival.get('time', '')
            row.cells[7].text = arrival.get('note', '')
        
        self.doc.add_paragraph()
    
    def _add_accommodation_table(self, hotels_data):
        """Add Accommodation table with room breakdown"""
        if not hotels_data:
            return
        
        # Section heading
        heading = self.doc.add_paragraph()
        run = heading.add_run('Accommodation')
        run.bold = True
        run.font.size = Pt(11)
        run.italic = True
        
        for hotel in hotels_data:
            # Hotel summary row
            table = self.doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            
            row = table.rows[0]
            row.cells[0].text = 'From'
            row.cells[1].text = 'To'
            row.cells[2].text = 'Hotel Name'
            row.cells[3].text = 'Meal'
            row.cells[4].text = 'Note'
            
            # Yellow headers
            for cell in row.cells:
                self._set_cell_shading(cell, 'FFFF00')
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Hotel data row
            data_row = table.add_row()
            data_row.cells[0].text = hotel.get('check_in', '')
            data_row.cells[1].text = hotel.get('check_out', '')
            data_row.cells[2].text = hotel.get('name', '')
            data_row.cells[3].text = hotel.get('board_basis', 'BB')
            data_row.cells[4].text = hotel.get('note', '')
            
            # Room breakdown table
            room_table = self.doc.add_table(rows=2, cols=6)
            room_table.style = 'Table Grid'
            
            # Room header
            room_header = room_table.rows[0]
            room_header.cells[0].text = 'Room Type'
            room_header.cells[1].text = 'SGL'
            room_header.cells[2].text = 'DBL'
            room_header.cells[3].text = 'TWIN'
            room_header.cells[4].text = 'TRPL'
            room_header.cells[5].text = 'OTHER'
            
            for cell in room_header.cells:
                self._set_cell_shading(cell, 'FFFF00')
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Room data
            room_data = room_table.rows[1]
            room_data.cells[0].text = 'Standard'
            room_data.cells[1].text = str(hotel.get('single_rooms', 0))
            room_data.cells[2].text = str(hotel.get('double_rooms', 0))
            room_data.cells[3].text = str(hotel.get('twin_rooms', 0))
            room_data.cells[4].text = str(hotel.get('triple_rooms', 0))
            room_data.cells[5].text = str(hotel.get('other_rooms', 0))
            
            # Notes row
            notes_row = room_table.add_row()
            notes_row.cells[0].text = 'Notes'
            notes_row.cells[0].merge(notes_row.cells[5])
            
            self.doc.add_paragraph()
        
        self.doc.add_paragraph()
    
    def _add_itinerary_table(self, itinerary_days):
        """Add Itinerary table"""
        if not itinerary_days:
            return
        
        # Section heading
        heading = self.doc.add_paragraph()
        run = heading.add_run('Itinerary')
        run.bold = True
        run.font.size = Pt(11)
        run.italic = True
        
        # Create table
        table = self.doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.autofit = False
        
        # Header row
        header_row = table.rows[0]
        header_row.cells[0].text = 'Date'
        header_row.cells[1].text = 'Itinerary Description'
        
        for cell in header_row.cells:
            self._set_cell_shading(cell, 'FFFF00')
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Set column widths
        table.columns[0].width = Inches(1.5)
        table.columns[1].width = Inches(5.5)
        
        # Add itinerary rows
        for day in itinerary_days:
            row = table.add_row()
            row.cells[0].text = day.get('date', '')
            row.cells[1].text = day.get('description', '')
        
        self.doc.add_paragraph()
    
    def _add_meals_table(self, meals_data):
        """Add Meals table"""
        if not meals_data:
            return
        
        # Section heading
        heading = self.doc.add_paragraph()
        run = heading.add_run('Meals')
        run.bold = True
        run.font.size = Pt(11)
        run.italic = True
        
        # Create table
        table = self.doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        # Header row
        headers = ['Date', 'Restaurant Name', 'Meal Type', 'Pax', 'Note']
        header_row = table.rows[0]
        for idx, header_text in enumerate(headers):
            cell = header_row.cells[idx]
            cell.text = header_text
            self._set_cell_shading(cell, 'FFFF00')
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add meal rows
        for meal in meals_data:
            row = table.add_row()
            row.cells[0].text = meal.get('date', '')
            row.cells[1].text = meal.get('restaurant', '')
            row.cells[2].text = meal.get('meal_type', '')
            row.cells[3].text = str(meal.get('pax', ''))
            row.cells[4].text = meal.get('note', '')
        
        self.doc.add_paragraph()
    
    def _add_transportation_table(self, transport_data):
        """Add Transportation table"""
        if not transport_data:
            return
        
        # Section heading
        heading = self.doc.add_paragraph()
        run = heading.add_run('Transportation')
        run.bold = True
        run.font.size = Pt(11)
        run.italic = True
        
        # Create table
        table = self.doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        # Header row
        headers = ['Time', 'Transportation Name', 'Note', 'Driver']
        header_row = table.rows[0]
        for idx, header_text in enumerate(headers):
            cell = header_row.cells[idx]
            cell.text = header_text
            self._set_cell_shading(cell, 'FFFF00')
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add transport rows
        for transport in transport_data:
            row = table.add_row()
            row.cells[0].text = transport.get('time', '')
            row.cells[1].text = transport.get('name', '')
            row.cells[2].text = transport.get('note', '')
            row.cells[3].text = transport.get('driver', '')
        
        self.doc.add_paragraph()
    
    def _add_guides_table(self, guides_data):
        """Add Guides table"""
        if not guides_data:
            return
        
        # Section heading
        heading = self.doc.add_paragraph()
        run = heading.add_run('Guides')
        run.bold = True
        run.font.size = Pt(11)
        run.italic = True
        
        # Create table with 5 columns
        table = self.doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        # Header row
        headers = ['From', 'To', 'Guide Name', 'Language', 'Note']
        header_row = table.rows[0]
        for idx in range(5):
            cell = header_row.cells[idx]
            cell.text = headers[idx]
            self._set_cell_shading(cell, 'FFFF00')
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add guide rows
        for guide in guides_data:
            row = table.add_row()
            row.cells[0].text = guide.get('from_date', '')
            row.cells[1].text = guide.get('to_date', '')
            row.cells[2].text = guide.get('name', '')
            row.cells[3].text = guide.get('language', '')
            row.cells[4].text = guide.get('note', '')
        
        self.doc.add_paragraph()
    
    def _add_miscellaneous_section(self):
        """Add Miscellaneous section"""
        heading = self.doc.add_paragraph()
        run = heading.add_run('Miscellaneous')
        run.bold = True
        run.font.size = Pt(11)
        run.italic = True
        
        self.doc.add_paragraph()
        
        # Page footer
        footer = self.doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = footer.add_run('Page 1 of 1')
        run.font.size = Pt(9)
    
    def generate_voucher(self, voucher_data, output_path=None):
        """
        Generate a trip voucher Word document in Windows of Jordan format
        
        Args:
            voucher_data (dict): Dictionary containing:
                - tour_file: Tour file number
                - company_name: Company name (default: Windows of Jordan)
                - tour: Tour info dict
                - arrivals: List of arrival/departure info
                - hotels: List of hotel details
                - itinerary_days: List of itinerary items
                - meals: List of meal details
                - transport: List of transport details
                - guides: List of guide details
            output_path (str): Optional path to save document
        
        Returns:
            str: Path to generated document
        """
        # Add header row
        self._add_header_row(
            tour_file=voucher_data.get('tour_file', voucher_data.get('voucher_number', 'DRAFT')),
            company_name=voucher_data.get('company_name', 'Windows of Jordan')
        )
        
        self.doc.add_paragraph()
        
        # Add tour details table
        if voucher_data.get('tour'):
            self._add_tour_details_table(voucher_data['tour'])
        
        # Add arrivals and departures
        if voucher_data.get('arrivals'):
            self._add_arrivals_departures_table(voucher_data['arrivals'])
        
        # Add accommodation
        if voucher_data.get('hotels'):
            self._add_accommodation_table(voucher_data['hotels'])
        
        # Add itinerary
        if voucher_data.get('itinerary_days'):
            self._add_itinerary_table(voucher_data['itinerary_days'])
        
        # Add meals
        if voucher_data.get('meals'):
            self._add_meals_table(voucher_data['meals'])
        
        # Add transportation
        if voucher_data.get('transport'):
            self._add_transportation_table(voucher_data['transport'])
        
        # Add guides
        if voucher_data.get('guides'):
            self._add_guides_table(voucher_data['guides'])
        
        # Add miscellaneous
        self._add_miscellaneous_section()
        
        # Save document
        if not output_path:
            output_dir = os.path.join(os.getcwd(), 'generated_vouchers')
            os.makedirs(output_dir, exist_ok=True)
            output_path = os.path.join(
                output_dir,
                f"Voucher_{voucher_data.get('tour_file', voucher_data.get('voucher_number', 'DRAFT'))}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            )
        
        self.doc.save(output_path)
        return output_path
