"""
Proforma Invoice Word Document Generator
Generates professional Word documents for proforma invoices with service line items
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import os


class ProformaDocGenerator:
    """Generate Word documents for proforma invoices"""
    
    def __init__(self):
        self.doc = Document()
        self._setup_styles()
    
    def _setup_styles(self):
        """Setup document styles"""
        # Set default font
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(10)
    
    def _add_header(self, company_name="Arabi Travel", company_address="Amman, Jordan"):
        """Add company header"""
        # Company name
        header = self.doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header.add_run(company_name)
        run.bold = True
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(139, 115, 85)  # Brand color
        
        # Company address
        address = self.doc.add_paragraph()
        address.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = address.add_run(company_address)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(107, 114, 128)
        
        # Add spacing
        self.doc.add_paragraph()
    
    def _add_title(self, invoice_number, invoice_date):
        """Add proforma invoice title"""
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run('PROFORMA INVOICE')
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(31, 41, 55)
        
        # Invoice details
        details = self.doc.add_paragraph()
        details.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = details.add_run(f'Invoice No: {invoice_number}')
        run.font.size = Pt(11)
        details.add_run('\n')
        run = details.add_run(f'Date: {invoice_date}')
        run.font.size = Pt(11)
        
        self.doc.add_paragraph()
    
    def _add_customer_info(self, customer_data):
        """Add customer information"""
        # Customer section
        customer_heading = self.doc.add_paragraph()
        run = customer_heading.add_run('Bill To:')
        run.bold = True
        run.font.size = Pt(12)
        
        # Customer details
        customer_info = self.doc.add_paragraph()
        customer_info.add_run(f"{customer_data.get('name', 'N/A')}\n").bold = True
        if customer_data.get('company_name'):
            customer_info.add_run(f"{customer_data['company_name']}\n")
        if customer_data.get('email'):
            customer_info.add_run(f"Email: {customer_data['email']}\n")
        if customer_data.get('phone'):
            customer_info.add_run(f"Phone: {customer_data['phone']}\n")
        if customer_data.get('nationality'):
            customer_info.add_run(f"Nationality: {customer_data['nationality']}\n")
        
        self.doc.add_paragraph()
    
    def _add_tour_details(self, tour_data):
        """Add tour information"""
        tour_heading = self.doc.add_paragraph()
        run = tour_heading.add_run('Tour Details:')
        run.bold = True
        run.font.size = Pt(12)
        
        # Tour info
        tour_info = self.doc.add_paragraph()
        if tour_data.get('from_date') and tour_data.get('to_date'):
            tour_info.add_run(f"Duration: {tour_data['from_date']} to {tour_data['to_date']}\n")
        if tour_data.get('pax'):
            tour_info.add_run(f"Number of Passengers: {tour_data['pax']}\n")
        if tour_data.get('nationality'):
            tour_info.add_run(f"Group Nationality: {tour_data['nationality']}\n")
        
        self.doc.add_paragraph()
    
    def _set_cell_border(self, cell, **kwargs):
        """Set border for table cell"""
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        
        # Create borders element
        tcBorders = OxmlElement('w:tcBorders')
        for edge in ('top', 'left', 'bottom', 'right'):
            edge_data = kwargs.get(edge, {'sz': 4, 'val': 'single', 'color': '#CCCCCC'})
            edge_el = OxmlElement(f'w:{edge}')
            for key, value in edge_data.items():
                edge_el.set(qn(f'w:{key}'), str(value))
            tcBorders.append(edge_el)
        
        tcPr.append(tcBorders)
    
    def _add_service_table(self, service_items):
        """Add table of service items"""
        # Create table with simple columns: #, Service, Pax, Total
        table = self.doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'
        table.autofit = False
        table.allow_autofit = False
        
        # Set column widths
        widths = [Inches(0.5), Inches(4.5), Inches(1.0), Inches(1.5)]
        for idx, width in enumerate(widths):
            for cell in table.columns[idx].cells:
                cell.width = width
        
        # Header row
        header_cells = table.rows[0].cells
        headers = ['#', 'Service Description', 'Passengers', 'Amount']
        
        for idx, header in enumerate(headers):
            cell = header_cells[idx]
            cell.text = header
            # Style header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(31, 41, 55)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Set background color
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'D1D5DB')
            cell._element.get_or_add_tcPr().append(shading_elm)
        
        # Add service rows
        for idx, item in enumerate(service_items, 1):
            row_cells = table.add_row().cells
            row_cells[0].text = str(idx)
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Service description
            row_cells[1].text = item.get('description', '')
            
            # Passengers
            row_cells[2].text = str(item.get('pax', ''))
            row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Amount
            row_cells[3].text = f"${item.get('total', 0):.2f}"
            row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Add totals row
        total_row = table.add_row().cells
        total_row[0].merge(total_row[2])
        total_row[0].text = 'TOTAL QUOTE'
        total_row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for paragraph in total_row[0].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(12)
        
        # Calculate and add total
        total_amount = sum(item.get('total', 0) for item in service_items)
        total_row[3].text = f"${total_amount:.2f}"
        total_row[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for paragraph in total_row[3].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(12)
        
        # Style total row
        for cell in [total_row[0], total_row[3]]:
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'FEF3C7')
            cell._element.get_or_add_tcPr().append(shading_elm)
        
        self.doc.add_paragraph()
    
    def _add_footer_notes(self):
        """Add footer with terms and conditions"""
        self.doc.add_paragraph()
        
        # Terms heading
        terms_heading = self.doc.add_paragraph()
        run = terms_heading.add_run('Terms & Conditions:')
        run.bold = True
        run.font.size = Pt(10)
        
        # Terms list
        terms = [
            'This is a proforma invoice and not a tax invoice.',
            'Payment is required before service delivery.',
            'All prices are in USD unless otherwise specified.',
            'Cancellation policy applies as per agreement.',
            'Please confirm this proforma to proceed with booking.'
        ]
        
        for term in terms:
            p = self.doc.add_paragraph(term, style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25)
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(107, 114, 128)
        
        self.doc.add_paragraph()
        
        # Thank you note
        thank_you = self.doc.add_paragraph()
        thank_you.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = thank_you.add_run('Thank you for your business!')
        run.italic = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(107, 114, 128)
    
    def generate_proforma(self, invoice_data, output_path=None):
        """
        Generate a proforma invoice Word document
        
        Args:
            invoice_data (dict): Dictionary containing:
                - invoice_number: Invoice number
                - invoice_date: Invoice date
                - customer: Customer info dict (name, company_name, email, phone, nationality)
                - tour: Tour info dict (from_date, to_date, pax, nationality)
                - service_items: List of service items with:
                    - description
                    - date_from
                    - date_to
                    - pax
                    - unit_price
                    - total
            output_path (str): Optional path to save document
        
        Returns:
            str: Path to generated document
        """
        # Add header
        self._add_header(
            company_name=invoice_data.get('company_name', 'Arabi Travel'),
            company_address=invoice_data.get('company_address', 'Amman, Jordan')
        )
        
        # Add title
        self._add_title(
            invoice_number=invoice_data.get('invoice_number', 'DRAFT'),
            invoice_date=invoice_data.get('invoice_date', datetime.now().strftime('%d %b %Y'))
        )
        
        # Add customer info
        if invoice_data.get('customer'):
            self._add_customer_info(invoice_data['customer'])
        
        # Add tour details
        if invoice_data.get('tour'):
            self._add_tour_details(invoice_data['tour'])
        
        # Add service items table
        if invoice_data.get('service_items'):
            self._add_service_table(invoice_data['service_items'])
        
        # Add footer
        self._add_footer_notes()
        
        # Save document
        if not output_path:
            # Create output directory if it doesn't exist
            output_dir = os.path.join(os.getcwd(), 'generated_invoices')
            os.makedirs(output_dir, exist_ok=True)
            output_path = os.path.join(
                output_dir,
                f"Proforma_{invoice_data.get('invoice_number', 'DRAFT')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            )
        
        self.doc.save(output_path)
        return output_path
